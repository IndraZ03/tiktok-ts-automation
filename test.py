from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

# Inisialisasi dokumen
doc = Document()

# Atur orientasi ke Landscape
section = doc.sections[-1]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = new_width
section.page_height = new_height

# Atur font default menjadi Arial 11
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Data Tabel Gabungan (Kegiatan 2 s.d 9)
data = [
    ["No", "Kegiatan", "Tahapan Kegiatan", "Output", "Keterkaitan dengan Agenda 2", "Kontribusi terhadap Visi Misi & Nilai-Nilai Organisasi", "Pihak Terkait", "Potensi Konflik", "Solusi Jika Ada Konflik"],
    
    # Kegiatan 2
    ["2", "Pengidentifikasian kebutuhan dan kondisi eksisting pelayanan data", "Menentukan aspek yang akan diidentifikasi seperti alur pelayanan, jenis layanan, fitur pelayanan, dan kebutuhan pengguna", "Tersusunnya aspek identifikasi kebutuhan dan kondisi eksisting pelayanan data", "Kompeten – Saya meningkatkan kualitas analisis dengan menentukan aspek identifikasi secara sistematis sebagai dasar pengembangan sistem pelayanan data.", "Berkontribusi terhadap peningkatan kualitas pelayanan BMKG melalui penyediaan sistem pelayanan data yang sesuai kebutuhan pengguna. Nilai organisasi: Profesional.", "Mentor, Ketua Tim Data dan Informasi", "Tidak – Tahapan dilakukan berdasarkan kebutuhan aktual pelayanan data.", "-"],
    ["", "", "Menyiapkan instrumen pengumpulan data berupa lembar observasi dan format inventarisasi dokumen", "Tersedianya instrumen pengumpulan data dan format inventarisasi dokumen", "Akuntabel – Saya menyiapkan instrumen pengumpulan data secara cermat dan dapat dipertanggungjawabkan untuk memperoleh data yang akurat.", "Mendukung tata kelola pelayanan yang efektif dan berbasis data. Nilai organisasi: Integritas.", "Mentor, Ketua Tim Data dan Informasi", "Tidak – Penyusunan instrumen dilakukan berdasarkan kebutuhan identifikasi.", "-"],
    ["", "", "Mengumpulkan dokumen pelayanan data, melakukan observasi, dan melakukan diskusi dengan pihak terkait", "Terkumpulnya dokumen pelayanan data dan hasil observasi kondisi eksisting pelayanan data", "Kolaboratif – Saya bekerja sama dengan pihak terkait dalam memperoleh data dan informasi yang dibutuhkan untuk pengembangan sistem pelayanan data.", "Mendukung peningkatan kualitas pelayanan melalui sinergi antarpegawai. Nilai organisasi: Sinergi.", "Ketua Tim Data dan Informasi, Petugas Pelayanan Data", "Ya – Keterbatasan waktu pegawai dan perbedaan informasi yang diperoleh dari masing-masing pihak.", "Melakukan koordinasi terlebih dahulu serta melakukan verifikasi silang terhadap data dan informasi yang diperoleh."],
    ["", "", "Merangkum hasil identifikasi dan menyusun kebutuhan sistem pelayanan data terpadu", "Tersusunnya dokumen kebutuhan sistem pelayanan data terpadu", "Adaptif – Saya menyusun kebutuhan sistem berdasarkan kondisi aktual dan kebutuhan pengguna sebagai dasar pengembangan layanan berbasis digital.", "Mendukung transformasi digital pelayanan BMKG yang lebih efektif dan berorientasi pada kebutuhan masyarakat. Nilai organisasi: Pelayanan.", "Mentor, Ketua Tim Data dan Informasi", "Ya – Perbedaan pendapat mengenai prioritas kebutuhan fitur yang akan dikembangkan.", "Menentukan prioritas kebutuhan berdasarkan urgensi, manfaat, dan hasil konsultasi dengan mentor serta pihak terkait."],
    
    # Kegiatan 3
    ["3", "Penyusunan rancangan sistem pelayanan data terpadu", "Menentukan kebutuhan fitur dan alur kerja sistem pelayanan data", "Tersusunnya daftar kebutuhan fitur dan alur kerja sistem pelayanan data", "Berorientasi Pelayanan – Saya mengidentifikasi kebutuhan fitur sistem berdasarkan kebutuhan pengguna layanan untuk memberikan kemudahan dan kepastian pelayanan data.", "Berkontribusi terhadap peningkatan kualitas pelayanan publik melalui pengembangan sistem pelayanan data yang sesuai kebutuhan masyarakat. Nilai organisasi: Pelayanan.", "Mentor, Ketua Tim Data dan Informasi, Petugas Pelayanan Data", "Ya – Perbedaan pendapat mengenai fitur yang menjadi prioritas untuk dikembangkan.", "Menentukan prioritas berdasarkan kebutuhan pengguna, urgensi masalah, dan hasil konsultasi dengan pihak terkait."],
    ["", "", "Menyiapkan perangkat lunak, referensi, dan dokumen pendukung perancangan sistem", "Tersedianya perangkat dan dokumen pendukung perancangan sistem", "Kompeten – Saya mempersiapkan perangkat dan referensi yang diperlukan untuk menghasilkan rancangan sistem yang berkualitas.", "Mendukung terwujudnya pelayanan yang profesional dan berbasis teknologi informasi. Nilai organisasi: Profesional.", "Peserta Aktualisasi", "Tidak – Tahapan dilakukan secara mandiri dengan sumber daya yang tersedia.", "-"],
    ["", "", "Menyusun workflow, desain antarmuka, dan rancangan fitur pelayanan data terpadu", "Tersusunnya workflow, desain antarmuka, dan rancangan fitur sistem pelayanan data terpadu", "Adaptif – Saya merancang sistem dengan memanfaatkan teknologi digital untuk meningkatkan efektivitas dan kualitas pelayanan data.", "Mendukung transformasi digital pelayanan BMKG yang lebih efektif dan modern. Nilai organisasi: Inovatif.", "Mentor, Ketua Tim Data dan Informasi", "Ya – Perbedaan masukan mengenai alur pelayanan dan desain sistem yang akan diterapkan.", "Mengakomodasi masukan yang relevan dan menyesuaikan rancangan dengan kebutuhan organisasi serta pengguna layanan."],
    ["", "", "Melakukan reviu dan penyempurnaan rancangan sistem berdasarkan masukan mentor dan pengguna", "Tersusunnya rancangan final sistem pelayanan data terpadu", "Kolaboratif – Saya bekerja sama dengan mentor dan pihak terkait dalam menyempurnakan rancangan sistem agar sesuai dengan kebutuhan organisasi.", "Mendukung peningkatan kualitas pelayanan melalui sinergi dan kerja sama antarpegawai. Nilai organisasi: Sinergi.", "Mentor, Ketua Tim Data dan Informasi, Petugas Pelayanan Data", "Ya – Perbedaan pendapat dalam menentukan rancangan akhir sistem.", "Mengutamakan musyawarah dan mempertimbangkan masukan berdasarkan manfaat serta kesesuaian dengan tujuan aktualisasi."],
    
    # Kegiatan 4
    ["4", "Pembuatan (Development) Sistem Pelayanan Data Terpadu", "Membangun basis data (database) sesuai rancangan kebutuhan", "Terbangunnya basis data sistem pelayanan", "Kompeten – Saya merancang arsitektur data secara teliti agar informasi tersimpan secara terstruktur.", "Mendukung tata kelola pelayanan yang efektif berbasis teknologi digital. Nilai organisasi: Profesional.", "Mentor, Ketua Tim Data dan Informasi", "Tidak – Pembangunan dilakukan secara teknis sesuai rancangan.", "-"],
    ["", "", "Melakukan pengkodean antarmuka (front-end) dan alur logika (back-end)", "Terbangunnya aplikasi sistem pelayanan data terpadu", "Adaptif – Saya mengimplementasikan teknologi dan bahasa pemrograman yang relevan dengan kebutuhan terkini.", "Mendukung transformasi digital pelayanan BMKG. Nilai organisasi: Inovatif.", "Mentor", "Ya – Munculnya kendala teknis (bug) pada saat pengkodean sistem.", "Mencari referensi penyelesaian masalah dan melakukan debugging pada kode sistem."],
    ["", "", "Mengintegrasikan basis data dengan sistem aplikasi", "Sistem berhasil terhubung dan dapat mengelola data", "Akuntabel – Saya memastikan jalur integrasi data aman dan terpercaya.", "Berkontribusi pada keandalan sistem pelayanan. Nilai organisasi: Integritas.", "Mentor", "Tidak – Integrasi dilakukan sesuai protokol standar.", "-"],
    
    # Kegiatan 5
    ["5", "Pelaksanaan Uji Coba (Testing) Sistem", "Melakukan uji coba fungsi internal (Alpha Testing)", "Catatan hasil uji coba fungsi internal (bug tracking)", "Kompeten – Saya memeriksa setiap fitur secara mendetail untuk memastikan tidak ada kesalahan fungsi.", "Menjamin mutu pelayanan informasi publik. Nilai organisasi: Profesional.", "Peserta Aktualisasi", "Tidak – Uji coba dilakukan secara mandiri.", "-"],
    ["", "", "Melakukan uji coba operasional (Beta Testing) bersama petugas pelayanan data", "Dokumen rekapitulasi masukan hasil uji coba operasional", "Kolaboratif – Saya melibatkan rekan kerja untuk mendapatkan masukan langsung dari sisi pengguna sistem.", "Meningkatkan sinergi antarpegawai dalam inovasi layanan. Nilai organisasi: Sinergi.", "Petugas Pelayanan Data", "Ya – Adanya fitur yang dianggap petugas kurang ramah pengguna (user-friendly).", "Mencatat setiap keluhan antarmuka dan menyesuaikannya agar lebih intuitif."],
    ["", "", "Melakukan perbaikan dan penyempurnaan sistem berdasarkan hasil uji coba", "Sistem pelayanan data terpadu versi final", "Berorientasi Pelayanan – Saya melakukan perbaikan demi kenyamanan pengguna dalam mengakses layanan.", "Memberikan layanan publik yang prima. Nilai organisasi: Pelayanan.", "Mentor", "Tidak – Perbaikan dilakukan murni berdasarkan data masukan.", "-"],
    
    # Kegiatan 6
    ["6", "Penyusunan Buku Panduan (Manual Book) Penggunaan Sistem", "Mengumpulkan bahan dan tangkapan layar (screenshot) sistem final", "Terkumpulnya aset visual panduan sistem", "Akuntabel – Saya mendokumentasikan antarmuka aplikasi sesuai dengan kondisi aktual yang siap rilis.", "Mendukung transparansi alur penggunaan sistem. Nilai organisasi: Integritas.", "Peserta Aktualisasi", "Tidak – Pengambilan bahan dilakukan secara mandiri.", "-"],
    ["", "", "Menyusun draf buku panduan penggunaan bagi admin dan pengguna", "Draf buku panduan (Manual Book)", "Berorientasi Pelayanan – Saya merangkai instruksi yang mudah dipahami demi kemudahan pengguna.", "Meningkatkan standar dan kepastian pelayanan publik. Nilai organisasi: Pelayanan.", "Mentor", "Tidak – Panduan disusun berdasarkan fitur final sistem.", "-"],
    ["", "", "Melakukan finalisasi panduan berdasarkan arahan mentor", "Dokumen final Buku Panduan Penggunaan Sistem", "Kolaboratif – Saya menyesuaikan bahasa teknis panduan agar selaras dengan masukan pimpinan.", "Menghasilkan panduan yang sesuai standar organisasi. Nilai organisasi: Sinergi.", "Mentor", "Ya – Perbedaan standar format penulisan dokumen organisasi.", "Mengikuti arahan format standar tata naskah yang berlaku di BMKG."],
    
    # Kegiatan 7
    ["7", "Pelaksanaan sosialisasi penggunaan sistem pelayanan data terpadu", "Menyusun materi dan metode sosialisasi penggunaan sistem", "Tersusunnya materi dan metode sosialisasi sistem pelayanan data terpadu", "Kompeten – Saya menyusun materi sosialisasi secara jelas dan sistematis agar mudah dipahami oleh peserta sosialisasi.", "Berkontribusi terhadap peningkatan kualitas pelayanan publik melalui penyebarluasan informasi penggunaan sistem pelayanan data. Nilai organisasi: Profesional.", "Mentor, Ketua Tim Data dan Informasi", "Tidak – Materi sosialisasi disusun berdasarkan fitur dan petunjuk teknis yang telah ditetapkan.", "-"],
    ["", "", "Menentukan peserta, jadwal, dan sarana pendukung sosialisasi", "Tersusunnya jadwal sosialisasi dan daftar peserta serta tersedianya sarana pendukung", "Harmonis – Saya menghargai kebutuhan dan ketersediaan waktu peserta dalam penentuan jadwal sosialisasi.", "Mendukung terciptanya lingkungan kerja yang kondusif dan kolaboratif. Nilai organisasi: Sinergi.", "Mentor, Ketua Tim Data dan Informasi, Pegawai terkait", "Ya – Keterbatasan waktu peserta dan perbedaan jadwal kegiatan masing-masing pihak.", "Melakukan koordinasi dan menentukan jadwal yang disepakati bersama."],
    ["", "", "Melaksanakan sosialisasi penggunaan sistem pelayanan data terpadu", "Terlaksananya sosialisasi penggunaan sistem pelayanan data terpadu", "Berorientasi Pelayanan – Saya memberikan penjelasan penggunaan sistem secara jelas untuk mendukung kemudahan akses layanan bagi pengguna dan pengelola layanan.", "Mendukung peningkatan kualitas pelayanan data yang lebih mudah diakses, transparan, dan efektif. Nilai organisasi: Pelayanan.", "Mentor, Ketua Tim Data dan Informasi, Pegawai terkait, Pengguna layanan", "Ya – Perbedaan tingkat pemahaman peserta terhadap penggunaan sistem yang disosialisasikan.", "Memberikan penjelasan secara bertahap, membuka sesi tanya jawab, dan menyediakan petunjuk teknis penggunaan sistem."],
    ["", "", "Mengumpulkan masukan dan melakukan evaluasi terhadap pelaksanaan sosialisasi", "Tersusunnya hasil evaluasi dan masukan terhadap pelaksanaan sosialisasi", "Kolaboratif – Saya menerima dan mengakomodasi masukan dari peserta sebagai bahan penyempurnaan sistem dan pelaksanaan layanan.", "Mendukung perbaikan berkelanjutan dalam penyelenggaraan pelayanan publik. Nilai organisasi: Integritas dan Sinergi.", "Mentor, Ketua Tim Data dan Informasi, Peserta Sosialisasi", "Ya – Adanya kritik atau masukan yang berbeda terkait penggunaan sistem.", "Menampung seluruh masukan secara objektif dan melakukan analisis untuk menentukan tindak lanjut yang diperlukan."],
    
    # Kegiatan 8
    ["8", "Pelaksanaan evaluasi terhadap efektivitas sistem dan efisiensi waktu pelayanan data", "Menentukan indikator efektivitas sistem dan efisiensi waktu pelayanan", "Tersusunnya indikator evaluasi efektivitas sistem dan efisiensi waktu pelayanan data", "Kompeten – Saya menyusun indikator evaluasi secara sistematis untuk mengukur keberhasilan implementasi sistem pelayanan data secara objektif.", "Berkontribusi terhadap peningkatan kualitas pelayanan publik melalui pengukuran kinerja layanan yang terarah dan terukur. Nilai organisasi: Profesional.", "Mentor, Ketua Tim Data dan Informasi", "Tidak – Indikator evaluasi disusun berdasarkan tujuan dan output kegiatan yang telah ditetapkan.", "-"],
    ["", "", "Menyiapkan instrumen evaluasi, data layanan, dan bukti pendukung", "Tersedianya instrumen evaluasi, data layanan, dan bukti pendukung evaluasi", "Akuntabel – Saya menyiapkan data dan instrumen evaluasi yang valid sebagai dasar penilaian hasil implementasi sistem.", "Mendukung tata kelola pelayanan yang transparan dan dapat dipertanggungjawabkan. Nilai organisasi: Integritas.", "Tim Data dan Informasi, Petugas Pelayanan Data", "Tidak – Tahapan dilakukan berdasarkan data dan dokumen yang tersedia.", "-"],
    ["", "", "Melaksanakan evaluasi dengan membandingkan kondisi sebelum dan sesudah implementasi sistem", "Terlaksananya evaluasi efektivitas sistem dan efisiensi waktu pelayanan data", "Akuntabel – Saya melakukan evaluasi secara objektif berdasarkan data dan fakta yang diperoleh dari pelaksanaan kegiatan.", "Mendukung peningkatan kualitas pelayanan data melalui proses evaluasi yang berkelanjutan. Nilai organisasi: Profesional.", "Mentor, Ketua Tim Data dan Informasi, Petugas Pelayanan Data", "Ya – Perbedaan persepsi dalam menilai tingkat keberhasilan implementasi sistem.", "Menggunakan indikator yang telah disepakati dan melakukan penilaian berdasarkan data yang terukur."],
    ["", "", "Menganalisis hasil evaluasi dan menyusun rekomendasi perbaikan berkelanjutan", "Tersusunnya laporan hasil evaluasi dan rekomendasi tindak lanjut", "Adaptif – Saya menyusun rekomendasi perbaikan berdasarkan hasil evaluasi untuk mendukung peningkatan kualitas pelayanan secara berkelanjutan.", "Mendukung transformasi digital dan peningkatan mutu pelayanan BMKG yang lebih efektif, efisien, dan responsif terhadap kebutuhan masyarakat. Nilai organisasi: Inovatif.", "Mentor, Ketua Tim Data dan Informasi", "Ya – Perbedaan pendapat mengenai prioritas rekomendasi perbaikan yang akan diterapkan.", "Menentukan rekomendasi berdasarkan tingkat urgensi, manfaat, dan kesepakatan bersama dengan pihak terkait."],
    
    # Kegiatan 9
    ["9", "Penyusunan laporan akhir aktualisasi", "Menentukan struktur dan sistematika laporan aktualisasi", "Tersusunnya struktur dan sistematika laporan aktualisasi", "Kompeten – Saya menyusun sistematika laporan secara sistematis agar hasil aktualisasi dapat terdokumentasikan dengan baik dan mudah dipahami.", "Berkontribusi terhadap peningkatan kualitas tata kelola organisasi melalui dokumentasi kegiatan yang tertib dan sistematis. Nilai organisasi: Profesional.", "Mentor, Coach", "Tidak – Tahapan dilakukan berdasarkan pedoman penyusunan laporan aktualisasi yang telah ditetapkan.", "-"],
    ["", "", "Mengumpulkan data, dokumen, dan eviden pelaksanaan aktualisasi", "Terkumpulnya data, dokumen, dan eviden pelaksanaan aktualisasi", "Akuntabel – Saya mengumpulkan eviden dan dokumen secara lengkap sebagai bentuk pertanggungjawaban pelaksanaan kegiatan aktualisasi.", "Mendukung terwujudnya tata kelola yang transparan dan dapat dipertanggungjawabkan. Nilai organisasi: Integritas.", "Mentor, Ketua Tim Data dan Informasi", "Ya – Terdapat dokumen atau eviden yang belum lengkap atau belum terdokumentasi dengan baik.", "Melakukan pengecekan kembali dokumen yang dibutuhkan dan berkoordinasi dengan pihak terkait untuk melengkapi eviden."],
    ["", "", "Menyusun laporan akhir aktualisasi berdasarkan hasil pelaksanaan kegiatan", "Tersusunnya draft laporan akhir aktualisasi", "Akuntabel – Saya menyusun laporan berdasarkan data dan hasil kegiatan yang sebenarnya sehingga dapat dipertanggungjawabkan.", "Mendukung penyelenggaraan organisasi yang transparan serta berorientasi pada hasil kerja. Nilai organisasi: Integritas.", "Mentor, Coach", "Ya – Adanya masukan atau koreksi terhadap isi laporan yang memerlukan penyesuaian.", "Menindaklanjuti masukan secara objektif dan melakukan revisi sesuai arahan mentor dan coach."],
    ["", "", "Melakukan reviu, perbaikan, dan finalisasi laporan aktualisasi", "Tersusunnya laporan akhir aktualisasi yang telah disetujui", "Kolaboratif – Saya bekerja sama dengan mentor dan coach dalam menyempurnakan laporan aktualisasi agar sesuai dengan ketentuan yang berlaku.", "Mendukung peningkatan kualitas hasil aktualisasi melalui sinergi dan kerja sama yang baik. Nilai organisasi: Sinergi.", "Mentor, Coach", "Ya – Perbedaan pendapat terkait substansi atau penyajian laporan aktualisasi.", "Melakukan pembahasan bersama dan menyesuaikan laporan berdasarkan arahan mentor, coach, serta ketentuan penyusunan laporan aktualisasi."]
]

# Pembuatan Tabel
table = doc.add_table(rows=1, cols=9)
table.style = 'Table Grid'

# Mengisi Header
hdr_cells = table.rows[0].cells
for i in range(9):
    hdr_cells[i].text = data[0][i]
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True

# Mengisi Baris Tabel
for row_data in data[1:]:
    row_cells = table.add_row().cells
    for i in range(9):
        row_cells[i].text = row_data[i]
        for paragraph in row_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)

# Simpan Dokumen
doc.save('Tabel_Kegiatan_Aktualisasi_Lengkap.docx')
print("File Tabel_Kegiatan_Aktualisasi_Lengkap.docx berhasil dibuat di direktori saat ini.")